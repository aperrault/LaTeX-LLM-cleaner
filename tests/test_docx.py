"""Tests for DOCX text extraction."""

import io
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document
from docx.shared import Inches, Pt

from latex_llm_cleaner.docx import extract_text_from_docx


def _create_tiny_png() -> bytes:
    """Create a minimal valid PNG image."""
    import struct
    import zlib

    width, height = 2, 2
    raw = b"\x00" + b"\xff\x00\x00" * width  # one scanline: filter byte + RGB
    raw *= height
    compressed = zlib.compress(raw)

    def chunk(ctype, data):
        c = ctype + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", compressed)
        + chunk(b"IEND", b"")
    )


def _create_test_docx(path: Path, paragraphs: list[str]) -> None:
    """Create a simple DOCX with paragraphs."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def _create_docx_with_headings(path: Path) -> None:
    """Create a DOCX with headings and paragraphs."""
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Intro text.")
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Section A content.")
    doc.add_heading("Subsection", level=3)
    doc.add_paragraph("Subsection content.")
    doc.save(str(path))


def _create_docx_with_table(path: Path) -> None:
    """Create a DOCX with a table."""
    doc = Document()
    doc.add_paragraph("Before table.")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "100"
    table.cell(2, 0).text = "Beta"
    table.cell(2, 1).text = "200"
    doc.add_paragraph("After table.")
    doc.save(str(path))


def _create_docx_with_image(path: Path) -> None:
    """Create a DOCX with an inline image."""
    doc = Document()
    doc.add_paragraph("Text before image.")
    png_bytes = _create_tiny_png()
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(2))
    doc.add_paragraph("Text after image.")
    doc.save(str(path))


def _create_docx_with_formatting(path: Path) -> None:
    """Create a DOCX with bold and italic text."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Normal text. ")
    bold_run = p.add_run("Bold text. ")
    bold_run.bold = True
    italic_run = p.add_run("Italic text.")
    italic_run.italic = True
    doc.save(str(path))


# --- Basic extraction tests ---


def test_basic_extraction(tmp_path):
    path = tmp_path / "test.docx"
    _create_test_docx(path, ["First paragraph.", "Second paragraph."])
    result = extract_text_from_docx(path)
    assert "First paragraph." in result
    assert "Second paragraph." in result


def test_heading_levels(tmp_path):
    path = tmp_path / "test.docx"
    _create_docx_with_headings(path)
    result = extract_text_from_docx(path)
    assert "# Title" in result
    assert "## Section A" in result
    assert "### Subsection" in result
    assert "Intro text." in result
    assert "Section A content." in result


def test_table_extraction(tmp_path):
    path = tmp_path / "test.docx"
    _create_docx_with_table(path)
    result = extract_text_from_docx(path)
    assert "| Name | Value |" in result
    assert "| Alpha | 100 |" in result
    assert "| Beta | 200 |" in result
    assert "| --- | --- |" in result
    assert "Before table." in result
    assert "After table." in result


def test_bold_italic_formatting(tmp_path):
    path = tmp_path / "test.docx"
    _create_docx_with_formatting(path)
    result = extract_text_from_docx(path)
    assert "Normal text." in result
    assert "**Bold text. **" in result
    assert "*Italic text.*" in result


# --- Image tests ---


def test_image_without_summary(tmp_path):
    path = tmp_path / "test.docx"
    _create_docx_with_image(path)
    result = extract_text_from_docx(path, verbose=False)
    assert "[Image]" in result
    assert "Text before image." in result
    assert "Text after image." in result


def test_image_with_summary(tmp_path):
    path = tmp_path / "test.docx"
    _create_docx_with_image(path)

    summary = tmp_path / "test_image1_summary.txt"
    summary.write_text("A tiny red and green pixel grid.")

    result = extract_text_from_docx(path)
    assert "[Image: A tiny red and green pixel grid.]" in result
    assert "[Image]" not in result


def test_custom_summary_suffix(tmp_path):
    path = tmp_path / "test.docx"
    _create_docx_with_image(path)

    summary = tmp_path / "test_image1_desc.md"
    summary.write_text("Custom description.")

    result = extract_text_from_docx(path, figure_summary_suffix="_desc.md")
    assert "[Image: Custom description.]" in result


# --- Comment tests ---


def test_notes_off_by_default(tmp_path):
    """Comments should not appear unless --notes is set."""
    path = Path("examples/Genesis_Internal_Proposal_v7.docx")
    if not path.exists():
        return  # skip if example not available
    result = extract_text_from_docx(path, notes=False)
    assert "Comment" not in result


def test_notes_includes_comments(tmp_path):
    """Comments should appear when --notes is set."""
    path = Path("examples/Genesis_Internal_Proposal_v7.docx")
    if not path.exists():
        return  # skip if example not available
    result = extract_text_from_docx(path, notes=True)
    assert "Comment (Alexandrov, Boian)" in result
    assert "Suggest replacing with printability" in result


# --- Auto-summarize DOCX tests ---


def test_auto_summarize_docx_generates_summaries(tmp_path):
    """Should generate summary files for DOCX images."""
    path = tmp_path / "doc.docx"
    _create_docx_with_image(path)

    mock_genai = MagicMock()
    mock_genai.Client.return_value = MagicMock()

    with patch(
        "latex_llm_cleaner.summarize._call_gemini_bytes",
        return_value="A small colored image.",
    ):
        import latex_llm_cleaner.summarize as mod

        with patch.dict(
            mod._run_batch_summarize.__globals__, {"genai": mock_genai}
        ):
            mod.auto_summarize_docx(path, {
                "google_api_key": "fake-key",
                "verbose": False,
                "figure_summary_suffix": "_summary.txt",
                "encoding": "utf-8",
            })

    summary_path = tmp_path / "doc_image1_summary.txt"
    assert summary_path.is_file()
    assert "colored image" in summary_path.read_text()


def test_auto_summarize_docx_skips_existing(tmp_path):
    """Should skip images that already have summary files."""
    path = tmp_path / "doc.docx"
    _create_docx_with_image(path)

    summary_path = tmp_path / "doc_image1_summary.txt"
    summary_path.write_text("Existing.")

    with patch(
        "latex_llm_cleaner.summarize._call_gemini_bytes",
    ) as mock_call:
        mock_genai = MagicMock()
        mock_genai.Client.return_value = MagicMock()

        import latex_llm_cleaner.summarize as mod

        with patch.dict(
            mod._run_batch_summarize.__globals__, {"genai": mock_genai}
        ):
            mod.auto_summarize_docx(path, {
                "google_api_key": "fake-key",
                "verbose": False,
                "figure_summary_suffix": "_summary.txt",
                "encoding": "utf-8",
            })

    mock_call.assert_not_called()
    assert summary_path.read_text() == "Existing."
